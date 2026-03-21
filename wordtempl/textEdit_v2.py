import sys
import json
import configparser
import os
from docx import Document

INI_PATH = r'C:/treeFrogProject/myapp/config.ini'

def mass_replace(config_path, section):
    print(f"\n=== Замена начата (Конфиг: {config_path}, Секция: {section}) ===")

    config = configparser.ConfigParser()
    if not config.read(config_path, encoding='utf-8'):
        print(f"Ошибка: Не удалось прочитать конфиг {config_path}")
        return

    if not config.has_section(section):
        print(f"Ошибка: Секция [{section}] не найдена в конфиге")
        return

    try:
        json_path = config.get(section, 'json')
        docx_path = config.get(section, 'copy')
    except Exception as e:
        print(f"Ошибка чтения путей из конфига: {e}")
        return

    if not os.path.exists(json_path):
        print(f"Ошибка: JSON файл не найден: {json_path}")
        return
    if not os.path.exists(docx_path):
        print(f"Ошибка: DOCX файл не найден: {docx_path}")
        return

    with open(json_path, 'r', encoding='utf-8') as f:
        replacements = json.load(f)

    doc = Document(docx_path)

    def replace_in_text_element(element):
        for para in element.paragraphs:
            replace_in_paragraph(para)

    def replace_in_paragraph(paragraph):
        full_text = ''.join([run.text for run in paragraph.runs])
        new_text = full_text
        for key, value in replacements.items():
            new_text = new_text.replace(str(key), str(value))

        if new_text != full_text:
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ''

    for para in doc.paragraphs:
        replace_in_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    doc.save(docx_path)
    print("=== Замена завершена ===\n")

# if __name__ == "__main__":
#     config_file = "C:/treeFrogProject/myapp/config.ini"
#     section_name = "TemplateReport"
#     # section_name = "TemplateBreed"
#     mass_replace(config_file, section_name)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: textEdit_v2.py <config_ini_path> <section_name>")
        sys.exit(1)

    config_file = sys.argv[1]
    section_name = sys.argv[2]
    # print(config_file, section_name)
    mass_replace(config_file, section_name)